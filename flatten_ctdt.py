import docx
import json
import re
import os

def find_table_by_keywords(doc, keywords, exclude_keywords=None):
    """Find a table containing all keywords and none of the exclude_keywords."""
    for i, table in enumerate(doc.tables):
        header_text = ""
        try:
            # Check first few rows for keywords to identify the table
            for tr in table._tbl.tr_lst[:5]:
                for tc in tr.tc_lst:
                    tc_text = "".join(node.text if node.text else "" for node in tc.iter() if node.tag.endswith('t'))
                    header_text += " " + clean_text(tc_text)
        except:
            continue
            
        header_text = header_text.lower()
        if all(k.lower() in header_text for k in keywords):
            if exclude_keywords and any(ek.lower() in header_text for ek in exclude_keywords):
                continue
            return table, i
    return None, -1

def get_raw_table_texts(table):
    """Extract text from a table using raw XML to bypass python-docx O(N^2) grid calculations."""
    rows = []
    for tr in table._tbl.tr_lst:
        row_cells = []
        for tc in tr.tc_lst:
            tc_text = "".join(node.text if node.text else "" for node in tc.iter() if node.tag.endswith('t'))
            row_cells.append(clean_text(tc_text))
        # Handle merged cells by extending the row to the max width found so far if needed, or just append
        rows.append(row_cells)
    
    # Pad rows to have the same length to avoid IndexError on data rows
    if rows:
        max_cols = max(len(r) for r in rows)
        for r in rows:
            r.extend([""] * (max_cols - len(r)))
            
    return rows

def clean_text(text):
    """Normalize text by removing extra spaces and newlines."""
    if not text: return ""
    return re.sub(r'\s+', ' ', text).strip()

def extract_course_list_12_2(table):
    """Extract course list from Table 12.2 (Khung chương trình)."""
    if not table: return []
    
    raw_rows = get_raw_table_texts(table)
    if not raw_rows: return []
    
    # Identify column roles
    header_rows_count = 1
    # Search for the row containing column names
    for r_idx in range(min(5, len(raw_rows))):
        row_text = " ".join([cell.lower() for cell in raw_rows[r_idx]])
        if "mã học phần" in row_text and "tên học phần" in row_text:
            header_rows_count = r_idx + 1
            break
            
    # Map column indices
    # We'll use the last header row to find column roles
    header_cells = [c.lower() for c in raw_rows[header_rows_count-1]]
    
    col_map = {
        "stt": -1, "ma_hp": -1, "ten_hp": -1, "so_tc": -1, 
        "lt": -1, "th": -1, "khac": -1, "tq": -1, "hk": -1
    }
    
    for idx, text in enumerate(header_cells):
        if "stt" in text or "tt" == text: col_map["stt"] = idx
        elif ("mã học phần" in text or "mã hp" in text) and "tiên quyết" not in text and "tiền đề" not in text and "trước" not in text: col_map["ma_hp"] = idx
        elif "tên học phần" in text or "tên hp" in text: col_map["ten_hp"] = idx
        elif "số tín chỉ" in text or "số tc" in text: col_map["so_tc"] = idx
        elif "lý thuyết" in text or "lt" == text: col_map["lt"] = idx
        elif "thực hành" in text or "th" == text: col_map["th"] = idx
        elif "khác" in text: col_map["khac"] = idx
        elif "tiên quyết" in text or "tq" in text or "tiền đề" in text or "trước" in text: col_map["tq"] = idx
        elif "học kỳ" in text or "hk" in text: col_map["hk"] = idx

    # If column roles weren't found in one row, search the whole header block
    for r in range(header_rows_count):
        row_cells = [c.lower() for c in raw_rows[r]]
        for idx, text in enumerate(row_cells):
            if col_map["stt"] == -1 and ("stt" in text or "tt" == text): col_map["stt"] = idx
            if col_map["ma_hp"] == -1 and ("mã học phần" in text or "mã hp" in text) and "tiên quyết" not in text and "tiền đề" not in text and "trước" not in text: col_map["ma_hp"] = idx
            if col_map["ten_hp"] == -1 and ("tên học phần" in text or "tên hp" in text): col_map["ten_hp"] = idx
            if col_map["so_tc"] == -1 and ("số tín chỉ" in text or "số tc" in text): col_map["so_tc"] = idx
            if col_map["lt"] == -1 and ("lý thuyết" in text or "lt" == text): col_map["lt"] = idx
            if col_map["th"] == -1 and ("thực hành" in text or "th" == text): col_map["th"] = idx
            if col_map["khac"] == -1 and "khác" in text: col_map["khac"] = idx
            if col_map["tq"] == -1 and ("tiên quyết" in text or "tq" in text or "tiền đề" in text or "trước" in text): col_map["tq"] = idx
            if col_map["hk"] == -1 and ("học kỳ" in text or "hk" in text): col_map["hk"] = idx

    results = []
    for r_idx in range(header_rows_count, len(raw_rows)):
        try:
            row = raw_rows[r_idx]
            ma_hp = row[col_map["ma_hp"]] if col_map["ma_hp"] != -1 else ""
            if not ma_hp or len(ma_hp) < 3: continue # Skip category headers
            
            course = {
                "STT": row[col_map["stt"]] if col_map["stt"] != -1 else "",
                "Ma_HP": ma_hp,
                "Ten_HP": row[col_map["ten_hp"]] if col_map["ten_hp"] != -1 else "",
                "So_TC": row[col_map["so_tc"]] if col_map["so_tc"] != -1 else "",
                "Ly_Thuyet": row[col_map["lt"]] if col_map["lt"] != -1 else "",
                "Thuc_Hanh": row[col_map["th"]] if col_map["th"] != -1 else "",
                "Khac": row[col_map["khac"]] if col_map["khac"] != -1 else "",
                "Mon_Tien_Quyet": row[col_map["tq"]] if col_map["tq"] != -1 else "",
                "Hoc_Ky": row[col_map["hk"]] if col_map["hk"] != -1 else ""
            }
            results.append(course)
        except: continue
        
    return results

def extract_mapping_15_3(table):
    """Extract mapping from Table 15.3 (CLO to PI/PLO)."""
    if not table: return []

    raw_rows = get_raw_table_texts(table)
    if not raw_rows: return []

    pi_row_idx, plo_row_idx, header_rows_count = -1, -1, 0
    
    # Scan for PI row (numbers like 1.1)
    for r_idx in range(min(10, len(raw_rows))):
        cells = [c.strip() for c in raw_rows[r_idx]]
        if any(re.search(r'\d+\.\d+', c) for c in cells):
            pi_row_idx = r_idx
            # Find PLO row above
            for prev_r in range(r_idx - 1, -1, -1):
                prev_cells = [c.strip().upper() for c in raw_rows[prev_r]]
                if any("PLO" in c for c in prev_cells):
                    plo_row_idx = prev_r
                    break
            if plo_row_idx == -1: plo_row_idx = max(0, r_idx - 1)
            header_rows_count = r_idx + 1
            break
            
    if pi_row_idx == -1: return []

    # Map column roles
    stt_col, hk_col, ten_hp_col, clo_col = 0, 1, 2, 3
    for r in range(header_rows_count):
        cells = [c.strip().lower() for c in raw_rows[r]]
        for idx, text in enumerate(cells):
            if "stt" in text or "tt" == text: stt_col = idx
            elif "học kỳ" in text or "hk" == text: hk_col = idx
            elif "tên học phần" in text or "tên hp" in text: ten_hp_col = idx
            elif "clo" in text: clo_col = idx

    col_mapping = {}
    pi_cells = raw_rows[pi_row_idx]
    plo_cells = raw_rows[plo_row_idx]
    
    for idx in range(clo_col + 1, len(pi_cells)):
        try:
            pi_val = pi_cells[idx]
            plo_val = plo_cells[idx]
            if not plo_val: # Handle merged
                for back_idx in range(idx - 1, clo_col, -1):
                    prev_plo = plo_cells[back_idx]
                    if prev_plo:
                        plo_val = prev_plo
                        break
            if pi_val and re.search(r'\d+\.\d+', pi_val):
                col_mapping[idx] = {"PLO": plo_val or "PLO?", "PI": pi_val}
        except: continue

    results = []
    current_course = None
    c_stt, c_hk, c_ten = "", "", ""
    
    for r_idx in range(header_rows_count, len(raw_rows)):
        try:
            cells = raw_rows[r_idx]
            stt = cells[stt_col]
            ten = cells[ten_hp_col]
            clo = cells[clo_col]
            
            if not clo or "tổng" in ten.lower(): continue
            
            if stt: c_stt = stt
            if ten: c_ten = ten
            hk = cells[hk_col]
            if hk: c_hk = hk
            
            if not current_course or (stt and stt != current_course.get("STT")):
                if current_course: results.append(current_course)
                current_course = {"STT": c_stt, "Ten_HP": c_ten, "Hoc_Ky": c_hk, "Mappings": []}
            
            for col_idx, meta in col_mapping.items():
                if col_idx < len(cells):
                    val = cells[col_idx]
                    if val and re.search(r'[234xX]', val):
                        level = int(val) if val.isdigit() else 3
                        current_course["Mappings"].append({
                            "CLO": clo, "PLO": meta["PLO"], "PI": meta["PI"], "Level": level
                        })
        except: continue

    if current_course: results.append(current_course)
    return results

def extract_ctdt_data(docx_path):
    """Main entry point for extracting all data from a CTĐT document."""
    print(f"Processing: {docx_path}")
    doc = docx.Document(docx_path)
    
    # 1. Extract Course List (12.2) - Support multiple tables and dummy tables
    courses = []
    keywords_courses = ["mã học phần", "tên học phần"]
    exclude_keywords = ["clo", "pi"]
    
    for table in doc.tables:
        header_text = ""
        try:
            for tr in table._tbl.tr_lst[:5]:
                for tc in tr.tc_lst:
                    tc_text = "".join(node.text if node.text else "" for node in tc.iter() if node.tag.endswith('t'))
                    header_text += " " + clean_text(tc_text)
        except: continue
        
        header_text = header_text.lower()
        if all(k.lower() in header_text for k in keywords_courses) and not any(ek.lower() in header_text for ek in exclude_keywords):
            table_courses = extract_course_list_12_2(table)
            for new_c in table_courses:
                if new_c["Ma_HP"] and not any(c["Ma_HP"] == new_c["Ma_HP"] for c in courses):
                    courses.append(new_c)

    
    # 2. Extract Mapping Matrix (15.3) - Support multiple tables if split
    matrix = []
    # Keywords to identify mapping tables
    keywords_matrix = ["clo", "pi", "học phần"]
    
    for table in doc.tables:
        header_text = ""
        try:
            for tr in table._tbl.tr_lst[:5]:
                for tc in tr.tc_lst:
                    tc_text = "".join(node.text if node.text else "" for node in tc.iter() if node.tag.endswith('t'))
                    header_text += " " + clean_text(tc_text)
        except: continue
        
        if all(k.lower() in header_text.lower() for k in keywords_matrix):
            table_data = extract_mapping_15_3(table)
            # Merge logic: if a course is split across tables, merge its mappings
            for new_course in table_data:
                existing = next((m for m in matrix if m["Ten_HP"] == new_course["Ten_HP"] or (m["STT"] == new_course["STT"] and m["STT"] != "")), None)
                if existing:
                    existing["Mappings"].extend(new_course["Mappings"])
                else:
                    matrix.append(new_course)
    
    # 3. Extract PLO/PI Text & Levels (Image 1)
    plo_list = []
    # Identify the PLO/PI table by keywords in header
    t_plo, _ = find_table_by_keywords(doc, ["ký hiệu", "csđg", "mô tả", "mức độ"], exclude_keywords=["học phần"])
    if not t_plo:
        t_plo, _ = find_table_by_keywords(doc, ["plo", "pi", "dự thảo", "đạt được"], exclude_keywords=["học phần"])
        
    if t_plo:
        current_plo = ""
        for row_idx in range(1, len(t_plo.rows)): # Skip header
            try:
                cells = [clean_text(c.text) for c in t_plo.rows[row_idx].cells]
                if len(cells) < 3: continue
                
                # Column 0: PLO code, Column 1: PI code, Column 2: Description, Column 3: Level
                plo_code = cells[0]
                pi_code = cells[1]
                desc = cells[2]
                level = cells[3] if len(cells) > 3 else ""
                
                if plo_code: current_plo = plo_code
                
                if current_plo or pi_code:
                    plo_list.append({
                        "Ma_PLO": current_plo,
                        "Ma_PI": pi_code,
                        "Mo_Ta": desc,
                        "Level": level
                    })
            except: continue

    return courses, matrix, plo_list

if __name__ == "__main__":
    import sys
    path = r"D:\Kiemdo\10.CTDT_He thong thong tin kinh doanh_HTTTQL_CQC.docx"
    if os.path.exists(path):
        c, m, p = extract_ctdt_data(path)
        print(f"Extracted {len(c)} courses, {len(m)} matrices, {len(p)} PLOs.")
        with open("final_extraction.json", "w", encoding="utf-8") as f:
            json.dump({"courses": c, "matrix": m, "plos": p}, f, ensure_ascii=False, indent=2)
