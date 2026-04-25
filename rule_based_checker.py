import docx
import json
import re
import os
from difflib import SequenceMatcher

class RuleBasedChecker:
    def __init__(self, db_path="db_mon_hoc.json"):
        self.db_path = db_path
        self.db_mon_hoc = []
        self.db_mon_hoc = []
        self.db_matrix = []
        self.db_plos = []
        self.load_db()

    def find_table_by_keywords(self, doc, keywords):
        for table in doc.tables:
            header_text = ""
            for r in range(min(3, len(table.rows))):
                header_text += " " + " ".join([cell.text for cell in table.rows[r].cells])
            header_text = header_text.lower()
            if all(k.lower() in header_text for k in keywords):
                return table
        return None

    def load_db(self):
        if os.path.exists(self.db_path):
            try:
                with open(self.db_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Support both old format (list) and new format (dict)
                    if isinstance(data, dict):
                        self.db_mon_hoc = data.get("courses", [])
                        self.db_matrix = data.get("matrix", [])
                        self.db_plos = data.get("plos", [])
                    else:
                        self.db_mon_hoc = data
            except Exception as e:
                print(f"Cảnh báo: Lỗi khi tải DB: {e}")

    def clean_text(self, text):
        if not text: return ""
        return re.sub(r'\s+', ' ', text).strip()

    def extract_full_text(self, doc):
        text = []
        for p in doc.paragraphs:
            text.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text)

    def find_value_after_keyword(self, text, keywords):
        for kw in keywords:
            # Match keyword followed by colon or space and then the value
            pattern = re.escape(kw) + r'[:\s]+([^\n\r]+)'
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return self.clean_text(match.group(1))
        return ""

    def check_syllabus(self, docx_path):
        print(f"Checking Syllabus: {docx_path}")
        try:
            doc = docx.Document(docx_path)
            full_text = self.extract_full_text(doc)
            errors = []

            # 1. Basic Form Check
            title_keywords = ["đề cương học phần", "đề cương môn học", "syllabus", "course specification"]
            if not any(kw in full_text.lower() for kw in title_keywords):
                errors.append("LỖI FORM: Không tìm thấy tiêu đề 'ĐỀ CƯƠNG HỌC PHẦN' hoặc 'SYLLABUS'.")

            # 2. Extract Key Info
            ma_hp = self.find_value_after_keyword(full_text, ["Mã số học phần", "Mã học phần", "Mã số môn học", "Course code"])
            ten_hp = self.find_value_after_keyword(full_text, ["Tên học phần", "Tên môn học", "Course name"])
            so_tc_str = self.find_value_after_keyword(full_text, ["Số tín chỉ", "Number of credits", "Credits"])
            
            # 3. Consistency with CTĐT
            db_course = None
            if ma_hp:
                db_course = next((c for c in self.db_mon_hoc if c.get("Ma_HP", "").lower() == ma_hp.lower()), None)
            
            if not db_course and ten_hp:
                # Try fuzzy match by name
                for c in self.db_mon_hoc:
                    ratio = SequenceMatcher(None, ten_hp.lower(), c.get("Ten_HP", "").lower()).ratio()
                    if ratio > 0.8:
                        db_course = c
                        break

            if db_course:
                # Check Code
                if ma_hp and db_course.get("Ma_HP", "").lower() != ma_hp.lower():
                    errors.append(f"LỖI ĐỐI SOÁT: Mã học phần '{ma_hp}' không khớp với CTĐT ({db_course.get('Ma_HP')}).")
                
                # Check Credits
                db_tc = str(db_course.get("So_TC", ""))
                if so_tc_str and db_tc and db_tc not in so_tc_str:
                    errors.append(f"LỖI ĐỐI SOÁT: Số tín chỉ '{so_tc_str}' không khớp với CTĐT ({db_tc}).")
                
                # Check Pre-requisites
                tq_syllabus = self.find_value_after_keyword(full_text, ["Học phần trước", "Môn học trước", "Prerequisites"])
                db_tq = db_course.get("Mon_Tien_Quyet", "")
                if db_tq and db_tq != "-" and db_tq.lower() not in tq_syllabus.lower():
                    errors.append(f"LỖI ĐỐI SOÁT: CTĐT yêu cầu học phần trước là '{db_tq}', nhưng đề cương ghi '{tq_syllabus}'.")

            # 4. Credit Distribution & Hours Calculation
            # Total = TC * 50
            # Theory = TC_LT * 15, Practice = TC_TH * 30
            match_tc = re.search(r'(\d+)\s*tín chỉ', so_tc_str) or re.search(r'(\d+)', so_tc_str)
            if match_tc:
                so_tc = int(match_tc.group(1))
                expected_total = so_tc * 50
                
                # Find hours in text
                total_hours_match = re.search(r'Phân bổ thời gian:\s*(\d+)\s*giờ', full_text, re.IGNORECASE)
                if total_hours_match:
                    actual_total = int(total_hours_match.group(1))
                    if actual_total != expected_total:
                        errors.append(f"LỖI QUY CHẾ: Tổng giờ học ({actual_total}) không khớp với số tín chỉ ({so_tc} TC * 50 = {expected_total} giờ).")

            # 5. Reference Year Check (>= 2020)
            ref_section_keywords = ["Tài liệu học tập", "Tài liệu bắt buộc", "Learning resources", "Textbooks"]
            for kw in ref_section_keywords:
                idx = full_text.lower().find(kw.lower())
                if idx != -1:
                    # Look at next 2000 characters
                    ref_text = full_text[idx:idx+2000]
                    years = re.findall(r'\b(19\d{2}|20[01]\d)\b', ref_text) # Years before 2020
                    if years:
                        # Only report if it looks like a bibliography entry
                        # We check if it's near the start of a line or after a dot
                        errors.append(f"CẢNH BÁO TÀI LIỆU: Phát hiện tài liệu xuất bản năm cũ ({', '.join(set(years))}). Quy định yêu cầu tài liệu từ 2020 trở lại đây.")
                        break

            # 6. Mapping Matrix Check (15.3 Consistency)
            # Find the mapping table in Syllabus
            db_course_matrix = None
            if db_course:
                db_course_matrix = next((m for m in self.db_matrix if (m.get("Ten_HP") and ten_hp and m.get("Ten_HP", "").lower() in ten_hp.lower()) or (m.get("STT") == db_course.get("STT"))), None)
            
            if db_course_matrix:
                # --- KIỂM TRA MỨC ĐỘ BLOOM (ẢNH 2) ---
                # 1. Tìm mức Bloom cao nhất yêu cầu trong CTĐT (15.3) cho môn này
                mappings = db_course_matrix.get("Mappings", [])
                if mappings:
                    max_bloom_required = max([m.get("Level", 0) for m in mappings if isinstance(m.get("Level"), int)] or [0])
                    
                    # 2. Tìm bảng B.1 (Ảnh 4) trong Syllabus để xem các bài đánh giá
                    t_assessment = self.find_table_by_keywords(doc, ["loại hình đánh giá", "phương pháp", "trọng số"])
                    if t_assessment:
                        # Giả định cột 3 là CLO (CĐR MH được đánh giá)
                        # Chúng ta cần AI hoặc Regex để bóc tách các bài đánh giá đạt mức nào.
                        # Tạm thời: Kiểm tra xem có bài nào có trọng số lớn (Cuối kỳ) đánh giá CLO đó không.
                        found_max_level = False
                        # Ở đây chúng ta cần bóc tách kỹ hơn ở Stage 3 (AI)
                        # Nhưng Logic Rule-based: Nếu Cuối kỳ đánh giá CLO quan trọng nhất thì tạm coi là đạt.
                        pass

                # --- ĐỐI SOÁT MA TRẬN 11.3 (ẢNH 3) ---
                t_11_3 = self.find_table_by_keywords(doc, ["ma trận tích hợp", "clo", "pi"])
                if t_11_3:
                    # So khớp từng cell của 11.3 với 15.3 trong DB
                    syllabus_mappings = []
                    # Bóc tách đơn giản 11.3
                    header_cells = [c.text.strip() for c in t_11_3.rows[0].cells]
                    for r_idx in range(1, len(t_11_3.rows)):
                        row = t_11_3.rows[r_idx].cells
                        clo = self.clean_text(row[0].text)
                        for c_idx in range(1, len(row)):
                            val = self.clean_text(row[c_idx].text)
                            if val and val.isdigit():
                                pi = header_cells[c_idx]
                                syllabus_mappings.append({"CLO": clo, "PI": pi, "Level": int(val)})
                    
                    # So sánh với DB
                    db_mappings = db_course_matrix.get("Mappings", [])
                    for db_m in db_mappings:
                        # Tìm mapping tương ứng trong syllabus
                        match = next((sm for sm in syllabus_mappings if sm["CLO"] == db_m["CLO"] and db_m["PI"] in sm["PI"]), None)
                        if not match:
                            errors.append(f"LỖI MA TRẬN: Bảng 11.3 thiếu mapping {db_m['CLO']} -> {db_m['PI']} như CTĐT.")
                        elif match["Level"] != db_m["Level"]:
                            errors.append(f"LỖI MA TRẬN: Mức độ {db_m['CLO']}->{db_m['PI']} là {match['Level']}, CTĐT yêu cầu {db_m['Level']}.")

                # --- KIỂM TRA NỘI DUNG PLO (ẢNH 1) ---
                for plo_item in self.db_plos:
                    ma_plo = plo_item.get("Ma_PLO")
                    if ma_plo and ma_plo in full_text:
                        # Kiểm tra xem mô tả có khớp nguyên văn không
                        db_desc = plo_item.get("Mo_Ta", "")
                        if db_desc and db_desc not in full_text:
                            # Fuzzy check
                            ratio = SequenceMatcher(None, db_desc.lower(), full_text.lower()).ratio()
                            if ratio < 0.5: # Rất khác
                                errors.append(f"CẢNH BÁO NỘI DUNG: Mô tả của {ma_plo} có vẻ không khớp nguyên văn với CTĐT.")

            return errors

        except Exception as e:
            return [f"LỖI HỆ THỐNG: {str(e)}"]

if __name__ == "__main__":
    checker = RuleBasedChecker("db_mon_hoc.json")
    # path = r"..."
    # errors = checker.check_syllabus(path)
    # print(errors)
